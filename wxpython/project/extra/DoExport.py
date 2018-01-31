# -*- coding: utf-8
__author__ = 'Sxn'
__date__ = '2017/5/24 12:32'


import os
import JsonIO
from JsonIO import JsonDict
import json
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

import sys
reload(sys)
sys.setdefaultencoding('ascii')


def exportJson():
    # if JsonIO.isConfirm[10] == False and JsonIO.cover['type'] == u'单位推荐':
    #     JsonIO.isConfirm[10] = True
    # print JsonIO.isConfirm
    COMFIRM = reduce(lambda x, y: x and y, JsonIO.isConfirm)
    # print COMFIRM
    if COMFIRM == False:
        return COMFIRM
        # pass

    # 创建放置Json文件和word的目录
    if os.path.isdir('ExportData') == False:
        os.mkdir('ExportData')
    if os.path.isdir('ExportWord') == False:
        os.mkdir('ExportWord')

    # 封面
    # if JsonIO.cover_dep_recommend['type'] == u'单位推荐':
    #     JsonIO.JsonDict['cover'] = JsonIO.cover_dep_recommend
    # else:
    #     JsonIO.JsonDict['cover'] = JsonIO.cover_expert_nomination
    JsonIO.JsonDict['cover'] = JsonIO.cover

    # 专家提名
    JsonIO.JsonDict['expert'] = JsonIO.expert

    # 申请人基本信息
    JsonIO.JsonDict['basic_info']['personal_info'] = JsonIO.personal_info
    JsonIO.JsonDict['basic_info']['dep_info'] = JsonIO.dep_info

    # 主要研究成果
    JsonIO.JsonDict['research_result'] = JsonIO.research_result

    # 拟展开的研究工作及其军事意义
    JsonIO.JsonDict['work_and_significance'] = JsonIO.work_and_significance

    # 学习经历
    JsonIO.JsonDict['personal_profile']['learning_exp'] = JsonIO.learning_exp

    # 工作经历
    JsonIO.JsonDict['personal_profile']['working_exp'] = JsonIO.working_exp

    # 承担国防相关代表性项目情况
    JsonIO.JsonDict['personal_profile']['projects_info'] = JsonIO.projects_info

    # 代表性成果
    JsonIO.JsonDict['personal_profile']['repr_results_to_word'] = JsonIO.repr_results_to_word

    # 重要科技奖项情况
    JsonIO.JsonDict['personal_profile']['awards'] = JsonIO.awards

    # 发明专利、国防专利情况
    JsonIO.JsonDict['personal_profile']['patent'] = JsonIO.patent

    JsonIO.outputName = u'国防科技卓越青年人才基金项目申报书_%s_%s' % (JsonDict['cover']['working_dep'], JsonDict['cover']['applicant'])

    JsonOutput = json.dumps(JsonIO.JsonDict)
    JsonOutFile = open('.\\ExportData\\' + JsonIO.outputName + '.json', 'w')
    JsonOutFile.write(JsonOutput)
    JsonOutFile.close()

    return True


def exportWord():
    isJsonFinish = exportJson()

    if isJsonFinish == False:
        return isJsonFinish

    # JsonDict = json.loads(open('test.json').read())

    doc = Document()

    font_HeiTi = u'黑体'
    font_SongTi = u'宋体'
    font_FangSong = u'仿宋_GB2312'

    # 标题字体
    style = doc.styles['Caption']
    style.font.bold = True
    style.font.size = Pt(26)

    # 段落标题字体
    style = doc.styles['Normal']
    style.font.bold = True
    style.font.size = Pt(14)
    # 段落标题括号内容字体
    style = doc.styles['Body Text 3']
    style.font.bold = False
    style.font.size = Pt(14)
    # 填写内容字体
    style = doc.styles['Body Text']     # paragraph
    style.font.bold = False
    style.font.size = Pt(12)

    style = doc.styles['Body Text Char']    # character
    style.font.bold = False
    style.font.size = Pt(12)
    # 段落中的注释字体
    style = doc.styles['Body Text 2']
    style.font.bold = False
    style.font.size = Pt(10)

    ''' 封面
    '''
    title = u'国防科技卓越青年人才基金\n项目申报书'
    doc.add_paragraph('\n\n')
    paraTitle = doc.add_paragraph()
    paraRun = paraTitle.add_run(title)
    setFont(paraRun, 26, True, font_HeiTi)
    doc.add_paragraph('\n' * 4)
    paraTitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    label1 = u'\t\t\t申报类型： '
    label2 = u'\t\t\t申 请 人： '
    label3 = u'\t\t\t项目名称： '
    label4 = u'\t\t\t申报领域： '
    label5 = u'\t\t\t工作单位： '
    label6 = u'\t\t\t联系电话： '

    type = doc.add_paragraph()
    typeTitleRun = type.add_run(label1)
    setFont(typeTitleRun, 16, True, font_HeiTi)
    typeRun = type.add_run(text=JsonDict['cover']['type'])
    setFont(typeRun, 16, False, font_SongTi)

    applicant = doc.add_paragraph()
    applicantTitleRun = applicant.add_run(label2)
    setFont(applicantTitleRun, 16, True, font_HeiTi)
    applicantRun = applicant.add_run(JsonDict['cover']['applicant'])
    setFont(applicantRun, 16, False, font_SongTi)

    project_name = doc.add_paragraph()
    project_nameTitleRun = project_name.add_run(label3)
    setFont(project_nameTitleRun, 16, True, font_HeiTi)
    project_nameRun = project_name.add_run(JsonDict['cover']['project_name'])
    setFont(project_nameRun, 16, False, font_SongTi)

    application_field = doc.add_paragraph()
    application_fieldTitleRun = application_field.add_run(label4)
    setFont(application_fieldTitleRun, 16, True, font_HeiTi)
    application_fieldRun = application_field.add_run(JsonDict['cover']['application_field'])
    setFont(application_fieldRun, 16, False, font_SongTi)

    working_dep = doc.add_paragraph()
    working_depTitleRun = working_dep.add_run(label5)
    setFont(working_depTitleRun, 16, True, font_HeiTi)
    working_depRun = working_dep.add_run(JsonDict['cover']['working_dep'])
    setFont(working_depRun, 16, False, font_SongTi)

    telephone = doc.add_paragraph()
    telephoneTitleRun = telephone.add_run(label6)
    setFont(telephoneTitleRun, 16, True, font_HeiTi)
    telephoneRun = telephone.add_run(JsonDict['cover']['telephone'])
    setFont(telephoneRun, 16, False, font_SongTi)

    bottom = doc.add_paragraph('\n' * 7)
    zyjw = bottom.add_run(u'中央军委科学技术委员会制\n')
    setFont(zyjw, 14, False, font_SongTi)
    finishDate = bottom.add_run(u'二〇一七年%s月%s日' % (' ' * 4, ' ' * 4))
    setFont(finishDate, 14, False, font_SongTi)
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_section(start_type=1)
    ''' 一、申请人基本信息
    '''
    baseInfo_title = doc.add_paragraph()
    baseInfo_titleRun = baseInfo_title.add_run(u'一、申请人基本信息')
    setFont(baseInfo_titleRun, 16, True, font_HeiTi)
    baseInfo = doc.add_table(rows=8, cols=7, style='Table Grid')

    # baseInfo.cell(2, 0).text = u'申请人情况'
    setInitCell(baseInfo.cell(2, 0), u'申请人情况', bold=True)
    baseInfo.cell(0, 0).merge(baseInfo.cell(1, 0)).merge(baseInfo.cell(2, 0)).merge(baseInfo.cell(3, 0)).merge(baseInfo.cell(4, 0))

    # baseInfo.cell(0, 1).text = u'姓 名'
    # baseInfo.cell(0, 2).text = JsonDict['basic_info']['personal_info']['name']
    # baseInfo.cell(0, 3).text = u'性 别'
    # baseInfo.cell(0, 4).text = JsonDict['basic_info']['personal_info']['gender']
    # baseInfo.cell(0, 5).text = u'出生年月'
    # baseInfo.cell(0, 6).text = JsonDict['basic_info']['personal_info']['birthday']
    setInitCell(baseInfo.cell(0, 1), u'姓 名')
    setEditCell(baseInfo.cell(0, 2), JsonDict['basic_info']['personal_info']['name'])
    setInitCell(baseInfo.cell(0, 3), u'性 别')
    setEditCell(baseInfo.cell(0, 4), JsonDict['basic_info']['personal_info']['gender'])
    setInitCell(baseInfo.cell(0, 5), u'出生年月')
    setEditCell(baseInfo.cell(0, 6), JsonDict['basic_info']['personal_info']['birthday'])

    # baseInfo.cell(1, 1).text = u'学 位'
    # baseInfo.cell(1, 2).text = JsonDict['basic_info']['personal_info']['bachelor']
    # baseInfo.cell(1, 3).text = u'职 称'
    # baseInfo.cell(1, 4).text = JsonDict['basic_info']['personal_info']['job']
    # baseInfo.cell(1, 5).text = u'单位职务'
    # baseInfo.cell(1, 6).text = JsonDict['basic_info']['personal_info']['dep_job']
    setInitCell(baseInfo.cell(1, 1), u'学 位')
    setEditCell(baseInfo.cell(1, 2), JsonDict['basic_info']['personal_info']['bachelor'])
    setInitCell(baseInfo.cell(1, 3), u'职 称')
    setEditCell(baseInfo.cell(1, 4), JsonDict['basic_info']['personal_info']['job'])
    setInitCell(baseInfo.cell(1, 5), u'单位职务')
    setEditCell(baseInfo.cell(1, 6), JsonDict['basic_info']['personal_info']['dep_job'])

    # baseInfo.cell(2, 1).text = u'主要研究方向'
    # baseInfo.cell(2, 2).text = JsonDict['basic_info']['personal_info']['research']
    setInitCell(baseInfo.cell(2, 1), u'主要研究方向')
    setEditCell(baseInfo.cell(2, 2), JsonDict['basic_info']['personal_info']['research'])
    baseInfo.cell(2, 2).merge(baseInfo.cell(2, 3)).merge(baseInfo.cell(2, 4)).merge(baseInfo.cell(2, 5)).merge(baseInfo.cell(2, 6))

    # baseInfo.cell(3, 1).text = u'身份证件名称'
    # baseInfo.cell(3, 2).text = JsonDict['basic_info']['personal_info']['id_name']
    # baseInfo.cell(3, 3).text = u'证件编号'
    # baseInfo.cell(3, 4).text = JsonDict['basic_info']['personal_info']['id_num']
    setInitCell(baseInfo.cell(3, 1), u'身份证件名称')
    setEditCell(baseInfo.cell(3, 2), JsonDict['basic_info']['personal_info']['id_name'])
    setInitCell(baseInfo.cell(3, 3), u'证件编号')
    setEditCell(baseInfo.cell(3, 4), JsonDict['basic_info']['personal_info']['id_num'])
    baseInfo.cell(3, 4).merge(baseInfo.cell(3, 5)).merge(baseInfo.cell(3, 6))

    # baseInfo.cell(4, 1).text = u'办公电话'
    # baseInfo.cell(4, 2).text = JsonDict['basic_info']['personal_info']['office_phone']
    # baseInfo.cell(4, 3).text = u'手 机'
    # baseInfo.cell(4, 4).text = JsonDict['basic_info']['personal_info']['telephone']
    # baseInfo.cell(4, 5).text = 'E-mail'
    # baseInfo.cell(4, 6).text = JsonDict['basic_info']['personal_info']['email']
    setInitCell(baseInfo.cell(4, 1), u'办公电话')
    setEditCell(baseInfo.cell(4, 2), JsonDict['basic_info']['personal_info']['office_phone'])
    setInitCell(baseInfo.cell(4, 3), u'手 机')
    setEditCell(baseInfo.cell(4, 4), JsonDict['basic_info']['personal_info']['telephone'])
    setInitCell(baseInfo.cell(4, 5), u'E-mail')
    setEditCell(baseInfo.cell(4, 6), JsonDict['basic_info']['personal_info']['email'])

    # baseInfo.cell(5, 0).text = u'申请单位情况'
    setInitCell(baseInfo.cell(5, 0), u'申请单位情况', bold=True)
    baseInfo.cell(5, 0).merge(baseInfo.cell(6, 0))

    # baseInfo.cell(5, 1).text = u'单位名称'
    # baseInfo.cell(5, 2).text = JsonDict['basic_info']['dep_info']['dep_name']
    setInitCell(baseInfo.cell(5, 1), u'单位名称')
    setEditCell(baseInfo.cell(5, 2), JsonDict['basic_info']['dep_info']['dep_name'])
    baseInfo.cell(5, 2).merge(baseInfo.cell(5, 3))
    # baseInfo.cell(5, 4).text = u'联系人（手机）'
    setInitCell(baseInfo.cell(5, 4), u'联系人（手机）')
    baseInfo.cell(5, 4).merge(baseInfo.cell(5, 5))
    # baseInfo.cell(5, 6).text = JsonDict['basic_info']['dep_info']['telephone']
    setEditCell(baseInfo.cell(5, 6), JsonDict['basic_info']['dep_info']['telephone'])

    # baseInfo.cell(6, 1).text = u'通信地址'
    setInitCell(baseInfo.cell(6, 1), u'通信地址')
    # baseInfo.cell(6, 2).text = JsonDict['basic_info']['dep_info']['address']
    setEditCell(baseInfo.cell(6, 2), JsonDict['basic_info']['dep_info']['address'])
    baseInfo.cell(6, 2).merge(baseInfo.cell(6, 3)).merge(baseInfo.cell(6, 4)).merge(baseInfo.cell(6, 5)).merge(baseInfo.cell(6, 6))

    text = u'''
        本人保证所填写的信息均真实有效，无任何虚假信息。本人完全清楚本声明的法律后果，如有不实，愿意承担相应的法律责任。\n
        %s项目申请人签字：
        %s\t\t年\t月\t日
    ''' % ('\t'*18, '\t'*18)
    # baseInfo.cell(7, 0).text = text
    setInitCell(baseInfo.cell(7, 0), text)
    # baseInfo.cell(7, 0).merge(baseInfo.cell(7, 1)).merge(baseInfo.cell(7, 2)).merge(baseInfo.cell(7, 3)).merge(baseInfo.cell(7, 4)).merge(baseInfo.cell(7, 5)).merge(baseInfo.cell(7, 6))

    reduce(lambda x, y: x.merge(y), [baseInfo.cell(7, i) for i in xrange(7)])

    doc.add_paragraph('\n')

    ''' 二、主要研究成果（2000字内）
    '''
    research_result_title = doc.add_paragraph()
    research_result_titleRun = research_result_title.add_run(u'二、主要研究成果')
    setFont(research_result_titleRun, 16, True, font_HeiTi)
    research_result_titleRun2 = research_result_title.add_run(u'（2000字内）')
    setFont(research_result_titleRun2, 16, False, font_FangSong)
    research_result_instruct = u'\t着重阐述近5年来在国防科技领域取得的代表性研究成果，如在国防基础研究领域取得重大发现，在工程研究领域突破重大技术瓶颈，在重要型号研制中攻克关键技术难题，或转化应用生产重大军事效益。'
    research_result_instructPara = doc.add_paragraph()
    research_result_instructRun = research_result_instructPara.add_run(research_result_instruct)
    setFont(research_result_instructRun, 14, False, font_FangSong)
    research_result = doc.add_paragraph()
    research_resultRun = research_result.add_run(JsonDict['research_result'] + '\n')
    setFont(research_resultRun, 12, False, font_SongTi)

    doc.add_paragraph('\n')

    ''' 三、拟展开的研究工作及其军事意义（2000字内）
    '''
    work_and_significance_title = doc.add_paragraph()
    work_and_significance_titleRun = work_and_significance_title.add_run(u'三、拟展开的研究工作及其军事意义')
    setFont(work_and_significance_titleRun, 16, True, font_HeiTi)
    work_and_significance_titleRun2 = work_and_significance_title.add_run(u'（2000字内）')
    setFont(work_and_significance_titleRun2, 16, False, font_FangSong)
    work_and_significance_instruct = u'\t着重阐述拟展开的研究工作的创新构思，主要研究方向和初步研究方案，及其军事意义和应用前景。'
    work_and_significance_instructPara = doc.add_paragraph()
    work_and_significance_instructRun = work_and_significance_instructPara.add_run(work_and_significance_instruct)
    setFont(work_and_significance_instructRun, 14, False, font_FangSong)
    research_result = doc.add_paragraph()
    research_resultRun = research_result.add_run(JsonDict['research_result'] + '\n')
    setFont(research_resultRun, 12, False, font_SongTi)

    doc.add_paragraph('\n')

    ''' 四、个人简介
    '''
    personal_profile = doc.add_paragraph()
    personal_profileRun = personal_profile.add_run(u'四、个人简介')
    setFont(personal_profileRun, 16, True, font_HeiTi)

    ''' （一）学习经历（从大学教育填起）
    '''
    learning_exp_title = doc.add_paragraph()
    learning_exp_titleRun = learning_exp_title.add_run(u'（一）学习经历')
    setFont(learning_exp_titleRun, 16, True, font_HeiTi)
    learning_exp_titleRun2 = learning_exp_title.add_run(u'（从大学教育填起）')
    setFont(learning_exp_titleRun2, 16, False, font_FangSong)

    learning_exp_number = len(JsonDict['personal_profile']['learning_exp'])
    learning_exp = doc.add_table(rows=learning_exp_number + 1, cols=4, style = 'Table Grid')

    # learning_exp.cell(0, 0).text = u'起止年月'
    # learning_exp.cell(0, 1).text = u'校（院）及系名称'
    # learning_exp.cell(0, 2).text = u'专业'
    # learning_exp.cell(0, 3).text = u'学位'
    setInitCell(learning_exp.cell(0, 0), u'起止年月')
    setInitCell(learning_exp.cell(0, 1), u'校（院）及系名称')
    setInitCell(learning_exp.cell(0, 2), u'专业')
    setInitCell(learning_exp.cell(0, 3), u'学位')

    for i in xrange(1, learning_exp_number + 1):
        # learning_exp.cell(i, 0).text = JsonDict['personal_profile']['learning_exp'][i-1]['start_end_date']
        # learning_exp.cell(i, 1).text = JsonDict['personal_profile']['learning_exp'][i-1]['school_name']
        # learning_exp.cell(i, 2).text = JsonDict['personal_profile']['learning_exp'][i-1]['major_subject']
        # learning_exp.cell(i, 3).text = JsonDict['personal_profile']['learning_exp'][i-1]['bachelor']
        setEditCell(learning_exp.cell(i, 0), JsonDict['personal_profile']['learning_exp'][i-1]['start_end_date'])
        setEditCell(learning_exp.cell(i, 1), JsonDict['personal_profile']['learning_exp'][i-1]['school_name'])
        setEditCell(learning_exp.cell(i, 2), JsonDict['personal_profile']['learning_exp'][i-1]['major_subject'])
        setEditCell(learning_exp.cell(i, 3), JsonDict['personal_profile']['learning_exp'][i-1]['bachelor'])

    doc.add_paragraph('\n')

    ''' （二）工作经历（含学术兼职情况）
    '''
    working_exp_title = doc.add_paragraph()
    working_exp_titleRun = working_exp_title.add_run(u'（二）工作经历')
    setFont(working_exp_titleRun, 16, True, font_HeiTi)
    working_exp_titleRun2 = working_exp_title.add_run(u'（含学术兼职情况）')
    setFont(working_exp_titleRun2, 16, False, font_FangSong)

    working_exp_number = len(JsonDict['personal_profile']['working_exp'])
    working_exp = doc.add_table(rows=working_exp_number + 1, cols=3, style='Table Grid')

    # working_exp.cell(0, 0).text = u'起止年月'
    # working_exp.cell(0, 1).text = u'工作单位'
    # working_exp.cell(0, 2).text = u'职务/职称'
    setInitCell(working_exp.cell(0, 0), u'起止年月')
    setInitCell(working_exp.cell(0, 1), u'工作单位')
    setInitCell(working_exp.cell(0, 2), u'职务/职称')

    for i in xrange(1, working_exp_number + 1):
        # working_exp.cell(i, 0).text = JsonDict['personal_profile']['working_exp'][i - 1]['start_end_date']
        # working_exp.cell(i, 1).text = JsonDict['personal_profile']['working_exp'][i - 1]['working_dep']
        # working_exp.cell(i, 2).text = JsonDict['personal_profile']['working_exp'][i - 1]['job']
        setEditCell(working_exp.cell(i, 0), JsonDict['personal_profile']['working_exp'][i - 1]['start_end_date'])
        setEditCell(working_exp.cell(i, 1), JsonDict['personal_profile']['working_exp'][i - 1]['working_dep'])
        setEditCell(working_exp.cell(i, 2), JsonDict['personal_profile']['working_exp'][i - 1]['job'])

    doc.add_paragraph('\n')

    ''' （三）承担国防相关代表性项目情况（限5项）
    '''
    projects_info_title = doc.add_paragraph()
    projects_info_titleRun = projects_info_title.add_run(u'（三）承担国防相关代表性项目情况')
    setFont(projects_info_titleRun, 16, True, font_HeiTi)
    projects_info_titleRun2 = projects_info_title.add_run(u'（限5项）')
    setFont(projects_info_titleRun2, 16, False, font_FangSong)

    projects_info_number = len(JsonDict['personal_profile']['projects_info'])
    projects_info = doc.add_table(rows=projects_info_number + 1, cols=8, style='Table Grid')

    # projects_info.cell(0, 0).text = u'序号'
    # projects_info.cell(0, 1).text = u'起止年月'
    # projects_info.cell(0, 2).text = u'名称'
    # projects_info.cell(0, 3).text = u'项目类别'
    # projects_info.cell(0, 4).text = u'来源'
    # projects_info.cell(0, 5).text = u'经费'
    # projects_info.cell(0, 6).text = u'本人承担任务'
    # projects_info.cell(0, 7).text = u'本人排序'
    setInitCell(projects_info.cell(0, 0), u'序号')
    setInitCell(projects_info.cell(0, 1), u'起止年月')
    setInitCell(projects_info.cell(0, 2), u'名称')
    setInitCell(projects_info.cell(0, 3), u'项目类别')
    setInitCell(projects_info.cell(0, 4), u'来源')
    setInitCell(projects_info.cell(0, 5), u'经费')
    setInitCell(projects_info.cell(0, 6), u'本人承担任务')
    setInitCell(projects_info.cell(0, 7), u'本人排序')

    for i in xrange(1, projects_info_number + 1):
        # projects_info.cell(i, 0).text = JsonDict['personal_profile']['projects_info'][i - 1]['id']
        # projects_info.cell(i, 1).text = JsonDict['personal_profile']['projects_info'][i - 1]['start_end_date']
        # projects_info.cell(i, 2).text = JsonDict['personal_profile']['projects_info'][i - 1]['project_name']
        # projects_info.cell(i, 3).text = JsonDict['personal_profile']['projects_info'][i - 1]['project_type']
        # projects_info.cell(i, 4).text = JsonDict['personal_profile']['projects_info'][i - 1]['source']
        # projects_info.cell(i, 5).text = JsonDict['personal_profile']['projects_info'][i - 1]['fund']
        # projects_info.cell(i, 6).text = JsonDict['personal_profile']['projects_info'][i - 1]['mission']
        # projects_info.cell(i, 7).text = JsonDict['personal_profile']['projects_info'][i - 1]['sort']
        setEditCell(projects_info.cell(i, 0), JsonDict['personal_profile']['projects_info'][i - 1]['id'])
        setEditCell(projects_info.cell(i, 1), JsonDict['personal_profile']['projects_info'][i - 1]['start_end_date'])
        setEditCell(projects_info.cell(i, 2), JsonDict['personal_profile']['projects_info'][i - 1]['project_name'])
        setEditCell(projects_info.cell(i, 3), JsonDict['personal_profile']['projects_info'][i - 1]['project_type'])
        setEditCell(projects_info.cell(i, 4), JsonDict['personal_profile']['projects_info'][i - 1]['source'])
        setEditCell(projects_info.cell(i, 5), JsonDict['personal_profile']['projects_info'][i - 1]['fund'])
        setEditCell(projects_info.cell(i, 6), JsonDict['personal_profile']['projects_info'][i - 1]['mission'])
        setEditCell(projects_info.cell(i, 7), JsonDict['personal_profile']['projects_info'][i - 1]['sort'])

    doc.add_paragraph('\n')

    ''' （四）代表性成果（论文、著作、研究技术报告、重要学术会议邀请报告，10篇以内第一作者或通信作者，按照重要性排序。每篇应说明申请人的主要贡献，包括：提出的学术思想、创造性、学术刊物中的主要引用及评价情况等。填写顺序：论文、著作、研究技术报告、重要学术会议邀请报告）
    '''
    repr_results_title = doc.add_paragraph()
    repr_results_titleRun = repr_results_title.add_run(u'（四）代表性成果')
    setFont(repr_results_titleRun, 16, True, font_HeiTi)
    repr_results_titleRun2 = repr_results_title.add_run(u'（论文、著作、研究技术报告、重要学术会议邀请报告，10篇以内第一作者或通信作者，按照重要性排序。每篇应说明申请人的主要贡献，包括：提出的学术思想、创造性、学术刊物中的主要引用及评价情况等。填写顺序：论文、著作、研究技术报告、重要学术会议邀请报告）')
    setFont(repr_results_titleRun2, 16, False, font_FangSong)

    repr_results_instructPara = doc.add_paragraph()
    repr_results_instructRun = repr_results_instructPara.add_run(u'填写顺序：\n')
    setFont(repr_results_instructRun, 14, False, font_FangSong)
    repr_results_instructRun = repr_results_instructPara.add_run(u' 论文：作者（按原排序），题目，期刊名称，卷（期）（年），起止页码；\n')
    setFont(repr_results_instructRun, 14, False, font_FangSong)
    repr_results_instructRun = repr_results_instructPara.add_run(u' 著作：作者（按原排序），著作名称，出版社，出版年份，出版地；\n')
    setFont(repr_results_instructRun, 14, False, font_FangSong)
    repr_results_instructRun = repr_results_instructPara.add_run(u' 研究技术报告（未公开发表的重要报告）：作者（按原排序），报告题目，完成年份；\n')
    setFont(repr_results_instructRun, 14, False, font_FangSong)
    repr_results_instructRun = repr_results_instructPara.add_run(u' 重要学术会议邀请报告：作者（按原排序），报告题目，报告年份，会议名称，地点；\n')
    setFont(repr_results_instructRun, 14, False, font_FangSong)

    repr_results_to_word_number = len(JsonDict['personal_profile']['repr_results_to_word'])
    repr_results_to_word = doc.add_table(rows=repr_results_to_word_number * 2 + 1, cols=5, style='Table Grid')
    # repr_results_to_word.cell(0, 0).text = u'序号'
    # repr_results_to_word.cell(0, 1).text = u'论著类型'
    # repr_results_to_word.cell(0, 2).text = u'作者（按原排序）'
    # repr_results_to_word.cell(0, 3).text = u'题目（名称）'
    # repr_results_to_word.cell(0, 4).text = u'论著相关信息'
    setInitCell(repr_results_to_word.cell(0, 0), u'序号')
    setInitCell(repr_results_to_word.cell(0, 1), u'论著类型')
    setInitCell(repr_results_to_word.cell(0, 2), u'作者（按原排序）')
    setInitCell(repr_results_to_word.cell(0, 3), u'题目（名称）')
    setInitCell(repr_results_to_word.cell(0, 4), u'论著相关信息')

    for i in xrange(1, repr_results_to_word_number + 1):
        # repr_results_to_word.cell(i * 2 - 1, 0).text = str(i)
        # repr_results_to_word.cell(i * 2 - 1, 1).text = JsonDict['personal_profile']['repr_results_to_word'][i - 1]['type']
        # repr_results_to_word.cell(i * 2 - 1, 2).text = JsonDict['personal_profile']['repr_results_to_word'][i - 1]['author']
        # repr_results_to_word.cell(i * 2 - 1, 3).text = JsonDict['personal_profile']['repr_results_to_word'][i - 1]['title']
        # repr_results_to_word.cell(i * 2 - 1, 4).text = JsonDict['personal_profile']['repr_results_to_word'][i - 1]['detail_info']
        # repr_results_to_word.cell(i * 2, 1).text = u'主要贡献及引用评价情况：' + JsonDict['personal_profile']['repr_results_to_word'][i - 1]['contribution_refs_info']
        setEditCell(repr_results_to_word.cell(i * 2 - 1, 0), str(i))
        setEditCell(repr_results_to_word.cell(i * 2 - 1, 1), JsonDict['personal_profile']['repr_results_to_word'][i - 1]['type'])
        setEditCell(repr_results_to_word.cell(i * 2 - 1, 2), JsonDict['personal_profile']['repr_results_to_word'][i - 1]['author'])
        setEditCell(repr_results_to_word.cell(i * 2 - 1, 3), JsonDict['personal_profile']['repr_results_to_word'][i - 1]['title'])
        setEditCell(repr_results_to_word.cell(i * 2 - 1, 4), JsonDict['personal_profile']['repr_results_to_word'][i - 1]['detail_info'])
        run = setInitCell(repr_results_to_word.cell(i * 2, 1), u'主要贡献及引用评价情况：', bold=True, alignCentre=False).add_run(JsonDict['personal_profile']['repr_results_to_word'][i - 1]['contribution_refs_info'])
        setFont(run, 12, False, font_SongTi)

        repr_results_to_word.cell(i * 2, 1).merge(repr_results_to_word.cell(i * 2, 2)).merge(repr_results_to_word.cell(i * 2, 3)).merge(repr_results_to_word.cell(i * 2, 4))

        repr_results_to_word.cell(i * 2 - 1, 0).merge(repr_results_to_word.cell(i * 2, 0))

    doc.add_paragraph('\n')

    ''' （五）重要科技奖项情况（10项以内）
    '''
    awards_title = doc.add_paragraph()
    awards_titleRun = awards_title.add_run(u'（五）重要科技奖项情况')
    setFont(awards_titleRun, 16, True, font_HeiTi)
    awards_titleRun2 = awards_title.add_run(u'（10项以内）')
    setFont(awards_titleRun2, 16, False, font_FangSong)
    awards_instruct = u'\t按顺序填写全部获奖人姓名（按原顺序）；获奖项目名称；获奖年份、类别及等级（如1999年国家自然科学二等奖，1998年军队科技进步一等奖等），并分别阐述申请人的主要贡献（限100字）。\n'
    awards_instructPara = doc.add_paragraph()
    awards_instructRun = awards_instructPara.add_run(awards_instruct)
    setFont(awards_instructRun, 14, False, font_FangSong)

    awards_number = len(JsonDict['personal_profile']['awards'])
    awards = doc.add_table(rows=awards_number * 2 + 1, cols=6, style='Table Grid')
    # awards.cell(0, 0).text = u'序号'
    # awards.cell(0, 1).text = u'全部获奖人'
    # awards.cell(0, 2).text = u'本人排序'
    # awards.cell(0, 3).text = u'获奖项目名称'
    # awards.cell(0, 4).text = u'获奖年份'
    # awards.cell(0, 5).text = u'获奖情况'
    setInitCell(awards.cell(0, 0), u'序号')
    setInitCell(awards.cell(0, 1), u'全部获奖人')
    setInitCell(awards.cell(0, 2), u'本人排序')
    setInitCell(awards.cell(0, 3), u'获奖项目名称')
    setInitCell(awards.cell(0, 4), u'获奖年份')
    setInitCell(awards.cell(0, 5), u'获奖情况')

    for i in xrange(1, awards_number + 1):
        # awards.cell(i * 2 - 1, 0).text = str(i)
        # awards.cell(i * 2 - 1, 1).text = JsonDict['personal_profile']['awards'][i - 1]['person_name']
        # awards.cell(i * 2 - 1, 2).text = JsonDict['personal_profile']['awards'][i - 1]['sort']
        # awards.cell(i * 2 - 1, 3).text = JsonDict['personal_profile']['awards'][i - 1]['project_name']
        # awards.cell(i * 2 - 1, 4).text = JsonDict['personal_profile']['awards'][i - 1]['year']
        # awards.cell(i * 2 - 1, 5).text = JsonDict['personal_profile']['awards'][i - 1]['category'] + JsonDict['personal_profile']['awards'][i - 1]['rank']
        # awards.cell(i * 2, 1).text = u'主要贡献：' + JsonDict['personal_profile']['awards'][i - 1]['contribution']
        setEditCell(awards.cell(i * 2 - 1, 0), str(i))
        setEditCell(awards.cell(i * 2 - 1, 1), JsonDict['personal_profile']['awards'][i - 1]['person_name'])
        setEditCell(awards.cell(i * 2 - 1, 2), JsonDict['personal_profile']['awards'][i - 1]['sort'])
        setEditCell(awards.cell(i * 2 - 1, 3), JsonDict['personal_profile']['awards'][i - 1]['project_name'])
        setEditCell(awards.cell(i * 2 - 1, 4), JsonDict['personal_profile']['awards'][i - 1]['year'])
        setEditCell(awards.cell(i * 2 - 1, 5), JsonDict['personal_profile']['awards'][i - 1]['category'] + JsonDict['personal_profile']['awards'][i - 1]['rank'])
        run = setInitCell(awards.cell(i * 2, 1), u'主要贡献：', bold=True, alignCentre=False).add_run(JsonDict['personal_profile']['awards'][i - 1]['contribution'])
        setFont(run, 12, False, font_SongTi)

        awards.cell(i * 2, 1).merge(awards.cell(i * 2, 2)).merge(awards.cell(i * 2, 3)).merge(awards.cell(i * 2, 4)).\
            merge(awards.cell(i * 2, 5))

        awards.cell(i * 2 - 1, 0).merge(awards.cell(i * 2, 0))

    doc.add_paragraph('\n')

    ''' （六）发明专利、国防专利情况（10项以内）
    '''
    patent_title = doc.add_paragraph()
    patent_titleRun = patent_title.add_run(u'（六）发明专利、国防专利情况')
    setFont(patent_titleRun, 16, True, font_HeiTi)
    patent_titleRun2 = patent_title.add_run(u'（10项以内）')
    setFont(patent_titleRun2, 16, False, font_FangSong)
    patent_instruct = u'\t请按顺序填写专利申报人（按原排序）；专利名称；申请年份、申请号；批准年份、专利号。并分别简述专利实施情况和申请人在专利发明和实施中的主要贡献（100字以内）。\n'
    patent_instructPara = doc.add_paragraph()
    patent_instructRun = patent_instructPara.add_run(patent_instruct)
    setFont(patent_instructRun, 14, False, font_FangSong)

    patent_number = len(JsonDict['personal_profile']['patent'])
    patent = doc.add_table(rows=patent_number * 2 + 1, cols=7, style='Table Grid')
    # patent.cell(0, 0).text = u'序号'
    # patent.cell(0, 1).text = u'专利申报人'
    # patent.cell(0, 2).text = u'专利名称'
    # patent.cell(0, 3).text = u'申请年份'
    # patent.cell(0, 4).text = u'申请号'
    # patent.cell(0, 5).text = u'批准年份'
    # patent.cell(0, 6).text = u'专利号'
    setInitCell(patent.cell(0, 0), u'序号')
    setInitCell(patent.cell(0, 1), u'专利申报人')
    setInitCell(patent.cell(0, 2), u'专利名称')
    setInitCell(patent.cell(0, 3), u'申请年份')
    setInitCell(patent.cell(0, 4), u'申请号')
    setInitCell(patent.cell(0, 5), u'批准年份')
    setInitCell(patent.cell(0, 6), u'专利号')

    for i in xrange(1, patent_number + 1):
        # patent.cell(i * 2 - 1, 0).text = str(i)
        # patent.cell(i * 2 - 1, 1).text = JsonDict['personal_profile']['patent'][i - 1]['person_name']
        # patent.cell(i * 2 - 1, 2).text = JsonDict['personal_profile']['patent'][i - 1]['patent_name']
        # patent.cell(i * 2 - 1, 3).text = JsonDict['personal_profile']['patent'][i - 1]['year']
        # patent.cell(i * 2 - 1, 4).text = JsonDict['personal_profile']['patent'][i - 1]['application_num']
        # patent.cell(i * 2 - 1, 5).text = JsonDict['personal_profile']['patent'][i - 1]['approved_year']
        # patent.cell(i * 2 - 1, 6).text = JsonDict['personal_profile']['patent'][i - 1]['patent_num']
        # patent.cell(i * 2, 1).text = u'主要贡献及专利实施情况：' + JsonDict['personal_profile']['patent'][i - 1]['contribution_conduct']
        setEditCell(patent.cell(i * 2 - 1, 0), str(i))
        setEditCell(patent.cell(i * 2 - 1, 1), JsonDict['personal_profile']['patent'][i - 1]['person_name'])
        setEditCell(patent.cell(i * 2 - 1, 2), JsonDict['personal_profile']['patent'][i - 1]['patent_name'])
        setEditCell(patent.cell(i * 2 - 1, 3), JsonDict['personal_profile']['patent'][i - 1]['year'])
        setEditCell(patent.cell(i * 2 - 1, 4), JsonDict['personal_profile']['patent'][i - 1]['application_num'])
        setEditCell(patent.cell(i * 2 - 1, 5), JsonDict['personal_profile']['patent'][i - 1]['approved_year'])
        setEditCell(patent.cell(i * 2 - 1, 6), JsonDict['personal_profile']['patent'][i - 1]['patent_num'])
        run = setInitCell(patent.cell(i * 2, 1), u'主要贡献及专利实施情况：', bold=True, alignCentre=False).add_run(JsonDict['personal_profile']['patent'][i - 1]['contribution_conduct'])
        setFont(run, 12, False, font_SongTi)

        patent.cell(i * 2, 1).merge(patent.cell(i * 2, 2)).merge(patent.cell(i * 2, 3)).merge(patent.cell(i * 2, 4)). \
            merge(patent.cell(i * 2, 5)).merge(patent.cell(i * 2, 6))

        patent.cell(i * 2 - 1, 0).merge(patent.cell(i * 2, 0))

    doc.add_paragraph('\n')

    ''' 五、单位审核
    '''
    check_title = doc.add_paragraph()
    check_titleRun = check_title.add_run(u'五、单位审核')
    setFont(check_titleRun, 16, True, font_HeiTi)

    check = doc.add_table(rows=2, cols=2, style='Table Grid')
    text = u'''
        %s单位负责人签字：
        %s  （单位公章）
        %s          \t年\t月\t日
    ''' % ('\t' * 7, '\t' * 7, '\t' * 7)
    # check.cell(0, 0).text = u'\n\t单位针对申报书真实性审核意见：' + '\n' * 12
    # check.cell(1, 1).text = text
    setInitCell(check.cell(0, 0), u'\t单位针对申报书真实性审核意见：' + '\n' * 12, alignCentre=False)
    setInitCell(check.cell(1, 1), text, alignCentre=False)

    newCell1 = check.cell(0, 0).merge(check.cell(0, 1))
    newCell2 = check.cell(1, 0).merge(check.cell(1, 1))
    newCell1.merge(newCell2)


    # doc.save('.\\ExportWord\\' + JsonIO.outputName + '.doc')
    doc.save('.\\ExportWord\\' + 'test' + '.doc')

    return True


def setFont(run, size, bold, font_name):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def setCell(cell, text, size, bold, font_name, alignCentre=True):
    para = cell.add_paragraph()
    run = para.add_run(text)
    setFont(run, size, bold, font_name)
    if alignCentre:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def setInitCell(cell, text, size=14, bold=False, font_name=u'仿宋_GB2312', alignCentre=True):
    return setCell(cell, text, size, bold, font_name, alignCentre)

def setEditCell(cell, text, size=12, bold=False, font_name=u'宋体', alignCentre=True):
    setCell(cell, text, size, bold, font_name, alignCentre)